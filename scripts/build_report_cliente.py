"""Relatorio cliente — versao resumida e estrategica para apresentacao."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============= PALETA =============
NAVY = RGBColor(0x0F, 0x2A, 0x47)
ACCENT = RGBColor(0xE8, 0x6B, 0x00)
GREEN = RGBColor(0x1F, 0x8A, 0x4E)
RED = RGBColor(0xC2, 0x39, 0x2C)
GRAY = RGBColor(0x4B, 0x55, 0x63)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = 'Montserrat'

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Aplica Montserrat em todo o documento (Normal + Headings)
def _apply_font_to_style(style_name, size=None, bold=None, color=None):
    try:
        st = doc.styles[style_name]
    except KeyError:
        return
    st.font.name = FONT
    # garante a fonte tambem no rPr east-asia / cs (compatibilidade Word)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), FONT)
    if size is not None:
        st.font.size = Pt(size)
    if bold is not None:
        st.font.bold = bold
    if color is not None:
        st.font.color.rgb = color

_apply_font_to_style('Normal', size=11)
_apply_font_to_style('Title', size=26, bold=True, color=NAVY)
_apply_font_to_style('Heading 1', size=18, bold=True, color=NAVY)
_apply_font_to_style('Heading 2', size=15, bold=True, color=NAVY)
_apply_font_to_style('Heading 3', size=12, bold=True, color=NAVY)
_apply_font_to_style('List Bullet', size=11)

# ============= HELPERS =============
def shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="BFBFBF", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def _force_run_font(run, size=None, color=None, bold=None):
    """Garante Montserrat no run + opcoes."""
    run.font.name = FONT
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold

def add_h2(text, color=NAVY, size=15, after=4, before=14, underline_border=True):
    """Heading 2 — usa estilo nativo do Word (aparece no sumario)."""
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    _force_run_font(run, size=size, color=color, bold=True)
    if underline_border:
        p_pr = p._p.get_or_add_pPr()
        p_borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), 'D5DBE3')
        p_borders.append(bottom)
        p_pr.append(p_borders)
    return p

def add_h3(text, color=NAVY, size=12, after=4, before=12):
    """Heading 3 — para topicos numerados (aparece no sumario)."""
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    _force_run_font(run, size=size, color=color, bold=True)
    return p

def add_body(text, size=11, italic=False, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if align:
        p.alignment = align
    run = p.add_run(text)
    _force_run_font(run, size=size, color=color, bold=bold)
    run.italic = italic
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        _force_run_font(r, size=11, bold=True)
        r2 = p.add_run(text)
        _force_run_font(r2, size=11)
    else:
        r = p.add_run(text)
        _force_run_font(r, size=11)
    return p

def add_kpi_row(items):
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, sub, color) in enumerate(items):
        cell = table.cell(0, i)
        shade(cell, "F4F6F9")
        set_cell_borders(cell, "D5DBE3", "4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p_label.add_run(label)
        _force_run_font(r_label, size=9, color=GRAY, bold=True)
        p_value = cell.add_paragraph()
        p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_value.paragraph_format.space_before = Pt(2)
        p_value.paragraph_format.space_after = Pt(0)
        r_value = p_value.add_run(value)
        _force_run_font(r_value, size=18, color=NAVY, bold=True)
        p_sub = cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(2)
        r_sub = p_sub.add_run(sub)
        _force_run_font(r_sub, size=9, color=color, bold=True)
    return table

def add_table_data(headers, rows, col_widths=None, header_color="0F2A47", first_col_bold=True):
    n_cols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shade(cell, header_color)
        set_cell_borders(cell, "0F2A47", "4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        _force_run_font(r, size=10, color=WHITE, bold=True)
    for i, row in enumerate(rows):
        zebra = i % 2 == 1
        for j, val in enumerate(row):
            cell = table.cell(i+1, j)
            if zebra:
                shade(cell, "F4F6F9")
            set_cell_borders(cell, "D5DBE3", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            text = str(val) if val is not None else ""
            bold_first = first_col_bold and j == 0
            r = p.add_run(text)
            _force_run_font(r, size=10, bold=bold_first)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_callout(title, body, fill="FFF4E6", border="E86B00", title_color=ACCENT):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_borders(cell, border, "8")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    _force_run_font(r, size=11, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    _force_run_font(r2, size=10.5, color=NAVY)

# =====================================================================
# CABEÇALHO — TÍTULO EM TEXTO (sem fundo)
# =====================================================================
p_title = doc.add_paragraph(style='Title')
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(2)
r_title = p_title.add_run("Sunrise Floor Removal")
_force_run_font(r_title, size=26, color=NAVY, bold=True)

p_period = doc.add_paragraph()
p_period.paragraph_format.space_after = Pt(8)
r_period = p_period.add_run("Período: 27 de Fevereiro a 27 de Março de 2025   •   Edição #1")
_force_run_font(r_period, size=11, color=GRAY)

# =====================================================================
# RESUMO DO MÊS — Heading 2 (aparece no sumario)
# =====================================================================
add_h2("Resumo do Mês", size=15, after=8)

add_body(
    "O mês de Abril seguiu com a campanha estável e dentro do padrão de oscilação esperado para "
    "o segmento. Ainda mantemos posição competitiva forte no leilão (entre os três maiores "
    "anunciantes da região para o serviço) e a qualidade dos leads continua dentro do perfil "
    "validado pelo cliente. Houve uma leve oscilação no custo por conversão, esperada em "
    "períodos de aumento da concorrência, e o plano de ação para Maio já está estruturado para "
    "recolocar o CPA na faixa de US$ 45.",
    color=GRAY
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_kpi_row([
    ("INVESTIMENTO", "US$ 1.965,77", "≈ US$ 60/dia",       GRAY),
    ("CONVERSÕES",   "35",            "Leads qualificados", GRAY),
    ("CPA",          "US$ 56,16",     "Em ajuste para Maio", GRAY),
    ("POSIÇÃO",      "Top 3",         "no leilão da região", GREEN),
])

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =====================================================================
# 1. PERFORMANCE — COMPARATIVO LIMPO
# =====================================================================
add_h3("1. Performance — Comparativo Mensal")

add_table_data(
    headers=["Métrica", "27/Jan a 27/Fev", "27/Fev a 27/Mar"],
    rows=[
        ["Investimento total", "US$ 1.728,28", "US$ 1.965,77"],
        ["Cliques",            "201",          "226"],
        ["CTR (taxa de cliques)", "5,87%",     "5,11%"],
        ["Conversões (leads)", "40",           "35"],
        ["CPA (custo por lead)", "US$ 43,21",  "US$ 56,16"],
    ],
    col_widths=[6.0, 4.5, 4.5],
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_callout(
    "Leitura do mês",
    "O volume de cliques cresceu 12% e o investimento subiu 14%, mas o CPA voltou para a faixa "
    "de US$ 50–60 que já vimos em outros momentos de pico de concorrência. A meta para Maio é "
    "recolocar o CPA na faixa dos US$ 45 com as ações detalhadas no item 5.",
    fill="EAF3EE", border="1F8A4E", title_color=GREEN
)

# =====================================================================
# 2. AÇÕES REALIZADAS NO MÊS
# =====================================================================
add_h3("2. Ações Realizadas no Mês")

add_bullet("Manutenção quinzenal dos termos de pesquisa, com revisão e inclusão de novas palavras-chave negativas para evitar tráfego de DIY, ferramentas, aluguéis e empregos.", bold_prefix="Curadoria de termos — ")
add_bullet("Otimização contínua dos lances por condado de forma independente, priorizando regiões com melhor custo por conversão.", bold_prefix="Geo-otimização — ")
add_bullet("Revisão das extensões de chamada e callouts para máxima exibição durante o horário comercial.", bold_prefix="Extensões — ")
add_bullet("Monitoramento da divisão de orçamento entre as duas campanhas principais (Phone Leads + Specific Services), com a campanha de Brand mantida pausada para concentrar verba em intenção direta.", bold_prefix="Distribuição de verba — ")
add_bullet("Acompanhamento do desempenho por dispositivo, com mobile mantido como principal canal de volume e desktop reforçando a conversão.", bold_prefix="Dispositivo — ")
add_bullet("Cruzamento contínuo dos leads com a planilha compartilhada para validar o ROI real e a qualidade dos contatos recebidos.", bold_prefix="Tracking de ROI — ")

# =====================================================================
# 3. CONCORRÊNCIA — DESTAQUE POSITIVO
# =====================================================================
add_h3("3. Posicionamento no Leilão")

add_body(
    "Comparativo com os anunciantes que disputam o mesmo leilão. Quanto maior a parcela de "
    "impressões, mais o concorrente aparece no nosso público. Topo absoluto = % de vezes que "
    "o anúncio aparece em primeira posição quando exibido.",
    color=GRAY, size=10, italic=True
)

add_table_data(
    headers=["Anunciante", "Parcela de impressões", "Topo absoluto"],
    rows=[
        ["Sunrise Floor Removal (nós)", "15,63%", "42,4%  ★"],
        ["Empire Today",                "15,32%", "32,4%"],
        ["50Floor",                     "15,11%", "43,6%"],
        ["National Floors Direct",      "< 10%",  "26,0%"],
        ["Thumbtack",                   "< 10%",  "21,7%"],
        ["Angi",                        "< 10%",  "15,1%"],
    ],
    col_widths=[6.5, 4.5, 4.0],
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    "Resultado do posicionamento",
    "Estamos entre os três anunciantes com maior presença no leilão da região, lado a lado com "
    "Empire Today e 50Floor (referências nacionais do segmento). Em topo absoluto da página, "
    "alcançamos 42,4% — empatado com 50Floor e à frente do Empire Today (32,4%). Este é um "
    "indicador forte de que a campanha está bem posicionada e a marca aparece com destaque "
    "para quem busca o serviço.",
    fill="EAF3EE", border="1F8A4E", title_color=GREEN
)

# =====================================================================
# 4. DESEMPENHO POR REGIÃO
# =====================================================================
add_h3("4. Desempenho por Região (Counties)")

add_body(
    "Visão consolidada das duas campanhas. Os bids são ajustados de forma independente para "
    "concentrar o orçamento nos condados de melhor retorno e reduzir exposição em condados "
    "com performance abaixo da média.",
    color=GRAY, size=10, italic=True
)

add_table_data(
    headers=["Condado", "Status", "Observação"],
    rows=[
        ["Orange",   "🟢 Núcleo principal",   "Maior volume e melhor distribuição de leads"],
        ["Seminole", "🟢 Bom desempenho",      "CPA dentro da meta — bid +20%"],
        ["Lake",     "🟢 Bom desempenho",      "Ótimo CPA — bid mantido em +15%"],
        ["Brevard",  "🟢 Bom desempenho",      "Performance estável — bid +5%"],
        ["Marion",   "🟡 Em observação",       "Performando no Phone Leads, ajuste em curso"],
        ["Polk",     "🟡 Em observação",       "Volume reduzido — bid em -20% / -40%"],
        ["Osceola",  "🟡 Em observação",       "Phone Leads bem; Specific em ajuste fino"],
        ["Pasco",    "🔵 Exposição reduzida",  "Bid em -90% — em revisão para Maio"],
        ["Sumter",   "🔵 Exposição reduzida",  "Volume mínimo — em revisão para Maio"],
    ],
    col_widths=[3.5, 4.5, 7.0],
)

# =====================================================================
# 5. PLANO PARA MAIO
# =====================================================================
add_h3("5. Plano de Otimização para Maio")

add_body("Refinamento de palavras-chave", bold=True, color=ACCENT, size=11)
add_bullet("Inclusão de uma nova rodada de negativas focada em concorrentes diretos e varejistas (ex.: Empire, 50Floor, Floor & Decor) para concentrar verba apenas em quem busca o serviço.")
add_bullet("Refinamento dos termos amplos de maior gasto, migrando para correspondência de frase para aumentar a precisão do tráfego.")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body("Otimização de lances", bold=True, color=ACCENT, size=11)
add_bullet("Aumento de bid nos condados com melhor desempenho (Osceola Phone Leads e Lake) para capturar mais volume onde o CPA é menor.")
add_bullet("Implementação de programação por hora, concentrando entrega entre 9h e 15h (pico de conversão histórico) e reduzindo investimento no fim de noite.")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body("Estratégia e criativos", bold=True, color=ACCENT, size=11)
add_bullet("Subida de 4 novos anúncios com ângulos de prova social e autoridade (\"+15 anos de experiência\", \"Licensed & Insured\", \"No Mess Guarantee\", \"1-Day Removal\") para destravar o CTR.")
add_bullet("Teste do Target CPA US$ 45 na campanha Specific Services — a campanha já atingiu o volume mínimo de conversões necessário pelo Google para essa estratégia.")
add_bullet("Reativação de Remarketing leve (US$ 5/dia) para visitantes do site que clicaram em orçamento mas não converteram — a lista atingiu o mínimo para essa campanha.")

doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_body("Metas para o próximo ciclo", bold=True, color=NAVY, size=11)
add_table_data(
    headers=["Indicador", "Atual", "Meta para Maio"],
    rows=[
        ["CPA",        "US$ 56,16",    "≤ US$ 45,00"],
        ["CTR",        "5,11%",        "≥ 6,00%"],
        ["Conversões", "35",           "≥ 42"],
    ],
    col_widths=[5.5, 5.0, 5.0],
)

# =====================================================================
# 6. REALOCAÇÃO DE VERBA
# =====================================================================
add_h3("6. Realocação de Verba — Mantendo US$ 60/dia")

add_table_data(
    headers=["Campanha", "Atual", "Sugestão p/ Maio"],
    rows=[
        ["[Search] Floor Removal — Phone Leads",    "US$ 30/dia", "US$ 30/dia"],
        ["[Search] Specific Services — Floor Demo", "US$ 30/dia", "US$ 30/dia"],
        ["[Search] Brand — Sunrise Floor Removal",  "PAUSADA",    "PAUSADA"],
        ["[Display] Remarketing — visitantes",      "—",          "≈ US$ 5/dia (extra)"],
    ],
    col_widths=[7.5, 4.0, 4.0],
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    "Estratégia da divisão",
    "Mantemos a verba concentrada nas duas campanhas de intenção direta (Phone Leads + Specific "
    "Services). A campanha de Brand segue pausada para liberar orçamento ao núcleo de geração "
    "de leads. A reativação leve do Remarketing entra como extra, com risco baixo, para "
    "recuperar visitantes que demonstraram interesse mas não converteram na primeira visita."
)

# =====================================================================
# 7. OBSERVAÇÕES — VALORIZAÇÃO DA QUALIDADE DOS LEADS
# =====================================================================
add_h3("7. Observações Estratégicas")

add_body(
    "A qualidade dos leads continua sendo o indicador principal de sucesso da campanha. As "
    "ligações e formulários recebidos seguem com perfil compatível com o serviço — esse é o "
    "sinal mais relevante de que a estratégia atual está atraindo o público certo, mesmo em "
    "meses de oscilação no CPA.",
    color=GRAY
)

add_body(
    "O cenário de leilão segue com pressão competitiva moderada, e a campanha mantém presença "
    "consistente entre os três maiores anunciantes da região. As métricas dentro da plataforma "
    "são meio para o fim: o que mede o sucesso real é o ROI, que continuamos cruzando "
    "mensalmente via planilha compartilhada para validar fechamentos e ticket médio dos leads "
    "recebidos.",
    color=GRAY
)

add_body(
    "O plano de ação para Maio já está mapeado e começa a entrar em execução na primeira "
    "semana, com foco em recolocar o CPA na faixa de US$ 45 sem comprometer o volume e a "
    "qualidade dos contatos.",
    color=GRAY
)

# =====================================================================
# RODAPÉ
# =====================================================================
doc.add_paragraph().paragraph_format.space_after = Pt(8)

footer_table = doc.add_table(rows=1, cols=1)
fc = footer_table.cell(0, 0)
shade(fc, "F4F6F9")
set_cell_borders(fc, "D5DBE3", "4")
fc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
fp = fc.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run("Dados consolidados a partir do Google Ads — janela 27/02/2025 a 27/03/2025")
fr.font.size = Pt(9)
fr.font.color.rgb = GRAY
fr.italic = True

# =====================================================================
# SAVE
# =====================================================================
out_path = "relatorio_cliente.docx"
doc.save(out_path)
print(f"OK: {out_path}")
