"""Generate the internal monthly .docx report - complete 10-section layout example."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============= PALETTE =============
NAVY = RGBColor(0x0F, 0x2A, 0x47)        # Navy titles
ACCENT = RGBColor(0xE8, 0x6B, 0x00)       # Orange accent
GREEN = RGBColor(0x1F, 0x8A, 0x4E)        # Positive green
RED = RGBColor(0xC2, 0x39, 0x2C)          # Negative red
GRAY = RGBColor(0x4B, 0x55, 0x63)         # Secondary text gray
LIGHT_GRAY = RGBColor(0xE8, 0xEC, 0xF1)   # Light background gray
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============= HELPERS =============
def shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="BFBFBF", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_h1(text, color=NAVY, size=22, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_h2(text, color=NAVY, size=14, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    # Subtle line below the title, implemented as a bottom border.
    p_pr = p._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'D5DBE3')
    p_borders.append(bottom)
    p_pr.append(p_borders)
    return p

def add_body(text, size=11, italic=False, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def add_kpi_row(items):
    """Create one row of KPI cards: items = [(label, value, delta, color), ...]."""
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value, delta, color) in enumerate(items):
        cell = table.cell(0, i)
        shade(cell, "F4F6F9")
        set_cell_borders(cell, "D5DBE3", "4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Reset the default paragraph spacing.
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        # label
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p_label.add_run(label)
        r_label.font.size = Pt(9)
        r_label.font.color.rgb = GRAY
        r_label.bold = True
        # value
        p_value = cell.add_paragraph()
        p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_value.paragraph_format.space_before = Pt(2)
        p_value.paragraph_format.space_after = Pt(0)
        r_value = p_value.add_run(value)
        r_value.font.size = Pt(18)
        r_value.bold = True
        r_value.font.color.rgb = NAVY
        # delta
        p_delta = cell.add_paragraph()
        p_delta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_delta.paragraph_format.space_before = Pt(0)
        p_delta.paragraph_format.space_after = Pt(2)
        r_delta = p_delta.add_run(delta)
        r_delta.font.size = Pt(9)
        r_delta.font.color.rgb = color
        r_delta.bold = True
    return table

def add_table_data(headers, rows, col_widths=None, header_color="0F2A47", first_col_bold=True):
    n_cols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shade(cell, header_color)
        set_cell_borders(cell, "0F2A47", "4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
    # Body
    for i, row in enumerate(rows):
        zebra = i % 2 == 1
        for j, val in enumerate(row):
            cell = table.cell(i+1, j)
            if zebra:
                shade(cell, "F4F6F9")
            set_cell_borders(cell, "D5DBE3", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            text = str(val) if val is not None else ""
            # color hint via prefix marker
            bold_first = first_col_bold and j == 0
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.bold = bold_first
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_callout(title, body, accent=ACCENT):
    """Callout box: one row with background color and text."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "FFF4E6")
    set_cell_borders(cell, "E86B00", "8")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = accent
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = NAVY

# =====================================================================
# HEADER / COVER
# =====================================================================

# Top band
header_table = doc.add_table(rows=1, cols=1)
header_cell = header_table.cell(0, 0)
shade(header_cell, "0F2A47")
header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = header_cell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("RELATÓRIO MENSAL  •  GOOGLE ADS")
r.font.color.rgb = WHITE
r.font.size = Pt(10)
r.bold = True
p2 = header_cell.add_paragraph()
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after = Pt(0)
r2 = p2.add_run("Sunrise Floor Removal")
r2.font.color.rgb = WHITE
r2.font.size = Pt(20)
r2.bold = True
p3 = header_cell.add_paragraph()
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(0)
r3 = p3.add_run("Período: 27 de Fevereiro a 27 de Março de 2025   •   Edição #1")
r3.font.color.rgb = RGBColor(0xC4, 0xD0, 0xDE)
r3.font.size = Pt(11)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =====================================================================
# EXECUTIVE SUMMARY
# =====================================================================
add_h1("Resumo Executivo", size=18, after=8)

add_body(
    "O mês fechou com investimento de US$ 1.965,77 e 35 conversões, gerando CPA de US$ 56,16. "
    "Em relação ao período anterior, houve aumento de 14% no investimento e queda de 13% nas conversões, "
    "puxados principalmente por (1) maior pressão competitiva no leilão (Empire Today e 50Floor "
    "subiram presença), (2) queda de CTR causada por termos amplos voltando a captar tráfego "
    "frio e (3) sazonalidade de fim de Abril. A tendência ainda é controlável, mas exige ação em "
    "negativas, programação por hora e bid por condado já no início de Maio.", color=GRAY
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_kpi_row([
    ("INVESTIMENTO",     "US$ 1.965,77",  "▲ +13,8% vs mês ant.",  RED),
    ("CONVERSÕES",       "35",            "▼ -12,5% vs mês ant.",  RED),
    ("CPA",              "US$ 56,16",     "▲ +30% vs mês ant.",    RED),
    ("CTR",              "5,11%",         "▼ -0,76 p.p.",          RED),
])

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =====================================================================
# 1. PERFORMANCE - COMPARISON TABLE
# =====================================================================
add_h2("1. Performance — Comparativo com o Mês Anterior")

add_table_data(
    headers=["Métrica", "27/Jan a 27/Fev", "27/Fev a 27/Mar", "Variação"],
    rows=[
        ["Investimento total",   "US$ 1.728,28", "US$ 1.965,77", "▲ +13,8%"],
        ["Cliques",              "201",          "226",          "▲ +12,4%"],
        ["Impressões",           "≈ 3.600",      "≈ 4.420",      "▲ +22,8%"],
        ["CTR",                  "5,87%",        "5,11%",        "▼ -0,76 p.p."],
        ["CPC médio",            "US$ 8,60",     "US$ 8,70",     "▲ +1,2%"],
        ["Taxa de conversão",    "19,90%",       "15,49%",       "▼ -4,4 p.p."],
        ["Conversões",           "40",           "35",           "▼ -12,5%"],
        ["CPA (custo/lead)",     "US$ 43,21",    "US$ 56,16",    "▲ +30,0%"],
    ],
    col_widths=[5.0, 3.8, 3.8, 3.0],
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    "Leitura técnica do mês",
    "Mais cliques + mais impressões + menos conversões = lead frio entrando no funil. "
    "O custo por clique ficou estável, ou seja, o leilão não encareceu — quem encareceu foram os "
    "termos amplos sem qualificação suficiente."
)

# =====================================================================
# 2. DIMENSION ANALYSIS (DEVICE + DAY)
# =====================================================================
add_h2("2. Análise por Dimensão")

add_body("Dispositivo — onde o investimento aterrissa", bold=True, size=11)
add_table_data(
    headers=["Dispositivo", "Investimento", "Cliques", "Conversões", "CPA"],
    rows=[
        ["Smartphones (mobile)", "US$ 1.371", "184", "≈ 22", "US$ 62,32"],
        ["Computadores (desktop)", "US$ 530", "55",  "≈ 11", "US$ 48,18"],
        ["Tablets",              "US$ 65",   "8",   "≈ 2",  "US$ 32,50"],
    ],
    col_widths=[5.0, 3.5, 3.0, 3.0, 3.0],
)
add_body(
    "Mobile segue como principal canal de volume (≈70% do gasto), porém neste mês desktop entregou "
    "CPA mais baixo (US$ 48 vs US$ 62). Esse padrão se inverte mês a mês: no acumulado dos 3 meses, "
    "mobile e desktop estão tecnicamente empatados (CVR 18,5% vs 17,7%). Manter mobile prioritário, "
    "sem ajuste agressivo.", italic=True, color=GRAY, size=10
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_body("Programação — quando os anúncios são vistos", bold=True, size=11)
add_table_data(
    headers=["Faixa de horário", "% das impressões", "Recomendação de bid"],
    rows=[
        ["00h – 06h (madrugada)", "0%",     "Cortar exibição (zerar) — atualmente sem schedule"],
        ["06h – 09h (manhã cedo)", "≈ 16%",  "+0% (manter)"],
        ["09h – 15h (pico)",      "≈ 50%",  "+10% (consolidar volume)"],
        ["15h – 18h (tarde)",     "≈ 18%",  "+0% (manter)"],
        ["18h – 24h (noite)",     "≈ 16%",  "−10% (reduzir gradualmente)"],
    ],
    col_widths=[4.5, 4.0, 7.5],
)

# =====================================================================
# 3. PERFORMANCE BY COUNTY
# =====================================================================
add_h2("3. Desempenho por Condado — Visão Consolidada")

add_table_data(
    headers=["Condado", "Investimento", "Conversões", "CPA", "Status"],
    rows=[
        ["Orange",   "US$ 782,73",  "15",    "US$ 52,18",  "✅ Núcleo principal"],
        ["Seminole", "US$ 215,57",  "5",     "US$ 43,11",  "✅ Bom desempenho"],
        ["Lake",     "US$ 251,83",  "5,5",   "US$ 45,79",  "✅ Bom desempenho"],
        ["Brevard",  "US$ 190,01",  "4",     "US$ 47,50",  "✅ Bom desempenho"],
        ["Marion",   "US$ 16,67",   "1",     "US$ 16,67",  "🟡 Phone Leads bem; Specific zerou"],
        ["Polk",     "US$ 77,34",   "2",     "US$ 38,67",  "🟡 Em recuperação"],
        ["Osceola",  "US$ 135,78",  "2",     "US$ 67,89",  "🟡 Specific elevou CPA"],
        ["Volusia",  "US$ 59,70",   "0",     "—",          "🚫 Pediu remoção em Fev — verificar"],
        ["Pasco",    "US$ 85,90",   "0",     "—",          "🚫 Pausar (sem conversão em 91 dias)"],
        ["Sumter",   "US$ 22,13",   "0",     "—",          "🚫 Pausar (sem conversão em 91 dias)"],
    ],
    col_widths=[3.5, 3.5, 2.8, 2.8, 5.0],
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    "Atenção — Volusia",
    "O condado de Volusia foi sinalizado pelo cliente para remoção em Fevereiro, mas seguiu recebendo "
    "exibições e gerando custo (US$ 59,70 no mês, US$ 284 no trimestre). Será verificado e removido "
    "definitivamente nas duas campanhas no início de Maio.",
    accent=RED
)

# =====================================================================
# 4. COMPETITION - AUCTION INSIGHTS
# =====================================================================
add_h2("4. Concorrência — Auction Insights")

add_body(
    "Análise de quem disputa o mesmo leilão. Quanto maior a parcela de impressões, "
    "mais o concorrente aparece no mesmo público que o nosso. Topo absoluto = % de vezes que "
    "aparece em primeira posição quando exibe.",
    color=GRAY, size=10, italic=True
)

add_table_data(
    headers=["Anunciante", "Parcela impr.", "Topo de página", "Topo absoluto"],
    rows=[
        ["Você (Sunrise Floor Removal)", "15,63%", "71,4%", "42,4%"],
        ["empiretoday.com",              "15,32%", "81,0%", "32,4%"],
        ["50floor.com",                  "15,11%", "80,0%", "43,6%"],
        ["nationalfloorsdirect.com",     "< 10%",  "65,1%", "26,0%"],
        ["thumbtack.com",                "< 10%",  "73,4%", "21,7%"],
        ["angi.com",                     "< 10%",  "74,7%", "15,1%"],
    ],
    col_widths=[5.5, 3.5, 3.5, 3.5],
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body(
    "Estamos competitivos: parcela de impressões equivalente aos dois maiores anunciantes do segmento "
    "(Empire Today e 50Floor) e a maior taxa de topo absoluto entre os concorrentes diretos do serviço "
    "(42,4%). Empire e 50Floor são os concorrentes que mais pressionam o leilão. Angi e Thumbtack são "
    "marketplaces, não competem pelo serviço final.",
    color=GRAY, size=10
)

# =====================================================================
# 5. SEARCH TERMS & KEYWORDS
# =====================================================================
add_h2("5. Termos de Pesquisa & Negativas")

add_body("Termos com melhor performance no mês (CPA < US$ 30)", bold=True, size=11)
add_table_data(
    headers=["Termo de pesquisa", "Cliques", "Conversões", "CPA"],
    rows=[
        ["miami floor removal",          "12", "3,0", "US$ 13,75"],
        ["floor removal miami",          "16", "4,0", "US$ 16,80"],
        ["tile removal miami",           "—",  "2,0", "US$ 9,62"],
        ["miami flooring removal",       "1",  "1,0", "US$ 9,65"],
        ["wood floor removal",             "11", "2,0", "US$ 33,98"],
        ["floor removal company near me",  "2",  "2,0", "US$ 25,01"],
        ["tile removal company",           "61", "9,0", "US$ 60,57"],
    ],
    col_widths=[6.5, 2.5, 2.8, 3.0],
)
add_body(
    "Padrão: termos com o nome da cidade e termos exatos de serviço convertem barato. "
    "Termos genéricos sem geo-modificador (ex.: 'carpet removal', 'dustless tile removal') entregam "
    "tráfego de pesquisa nacional/curiosidade.",
    color=GRAY, size=10, italic=True
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body("Termos que consumiram verba sem converter (mês)", bold=True, size=11)
add_table_data(
    headers=["Termo de pesquisa", "Custo", "Cliques", "Ação"],
    rows=[
        ["dustless tile removal (ampla)",  "US$ 122,99", "8",  "Refinar para frase/exata"],
        ["tile removal miami (ampla)",   "US$ 65,67",  "5",  "Refinar para frase"],
        ["carpet removal (ampla)",         "US$ 57,23",  "4",  "Refinar para frase"],
        ["empire flooring near me",        "US$ 26,85",  "1",  "Negativar (concorrente)"],
        ["floor and decor (variações)",    "≈ US$ 32",   "8",  "Negativar (varejista)"],
        ["dustram, dust ram",              "US$ 37,38",  "5",  "Negativar (concorrente)"],
        ["floor busters",                  "US$ 19,08",  "2",  "Negativar (concorrente local)"],
    ],
    col_widths=[6.5, 2.5, 2.0, 3.8],
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    "🚫 Novas negativas a aplicar (22 termos)",
    "Concorrentes/marcas: empire, empire today, empire flooring, 50floor, angi, angie, thumbtack, "
    "creative floor(s), floor busters, the floor fathers, bottoms up floor, stanley steemer, "
    "coronado carpet, dustram, dust ram, speedy floor, nationalfloorsdirect.\n"
    "Varejo / off-topic: floor and decor, floor & decor, melbourne, refinish, sanding, restoration."
)

# =====================================================================
# 6. WORK COMPLETED THIS MONTH
# =====================================================================
add_h2("6. O Que Foi Feito no Mês")

add_bullet("Manutenção quinzenal dos termos de pesquisa, com inclusão pontual de negativas em variações de DIY/aluguel.", bold_prefix="Negativas — ")
add_bullet("Continuidade do ajuste por condado (Pasco e Sumter mantidos em −90% / −20%).", bold_prefix="Geo — ")
add_bullet("Manutenção do split atual (US$ 30/dia Phone Leads + US$ 30/dia Specific Services). Brand pausada.", bold_prefix="Verba — ")
add_bullet("Revisão de extensões de chamada e callouts ativos durante o horário comercial.", bold_prefix="Extensões — ")
add_bullet("Monitoramento contínuo de leads via planilha de cruzamento com o cliente para validar ROI real.", bold_prefix="Tracking — ")

# =====================================================================
# 7. WHAT DROVE THE RESULT
# =====================================================================
add_h2("7. O Que Puxou o Resultado")

add_bullet("Termos amplos voltando a captar tráfego frio (`dustless tile removal`, `carpet removal`, `tile removal miami`).", bold_prefix="Tráfego frio — ")
add_bullet("Empire Today e 50Floor com presença ainda forte no leilão (parcela ≈15% cada, próxima da nossa).", bold_prefix="Pressão competitiva — ")
add_bullet("Conversões diluídas em condados secundários (Osceola, Polk e Pasco com Specific Services entregando CPA acima de US$ 65).", bold_prefix="Distribuição geográfica — ")
add_bullet("Volusia ainda gerando custo após pedido de remoção do cliente em Fevereiro.", bold_prefix="Falha pontual — ")

# =====================================================================
# 8. ACTION PLAN - MAY
# =====================================================================
add_h2("8. Plano de Ação para Maio")

add_body("Quick wins (semana 1)", bold=True, color=ACCENT, size=11)
add_bullet("Aplicar 22 novas negativas (concorrentes + varejo + cidades fora do escopo).", bold_prefix="🚫 ")
add_bullet("Pausar Pasco e Sumter nas duas campanhas (0 conv em 91 dias, US$ 234 e US$ 112 gastos respectivamente).", bold_prefix="📍 ")
add_bullet("Confirmar e remover Volusia (sinalizado pelo cliente em Fev mas seguiu rodando).", bold_prefix="📍 ")
add_bullet("Refinar `dustless tile removal` e `tile removal miami` ampla para correspondência de frase.", bold_prefix="🔍 ")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body("Otimização de bid (semanas 2–3)", bold=True, color=ACCENT, size=11)
add_bullet("Subir Osceola Phone Leads (CPA US$ 26,65) de −10% para 0%.", bold_prefix="↗ ")
add_bullet("Subir Lake Phone Leads (CPA US$ 31) de +15% para +25%.", bold_prefix="↗ ")
add_bullet("Reduzir Seminole no Specific Services (CPA US$ 62) de −15% para −30%.", bold_prefix="↘ ")
add_bullet("Aplicar programação por hora: −10% das 18h–22h, +10% das 9h–15h, cortar 22h–6h.", bold_prefix="🕐 ")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body("Estrutura e estratégia (semanas 3–4)", bold=True, color=ACCENT, size=11)
add_bullet("Subir 4 novas variações de anúncio com prova social + autoridade ('15+ years of service', 'Licensed & Insured', 'No Mess Guarantee', '1-Day Removal').", bold_prefix="✏ ")
add_bullet("Testar Target CPA US$ 45 no Specific Services (já tem 30+ conv, atende mínimo do Google).", bold_prefix="🎯 ")
add_bullet("Reativar lista de remarketing (visitantes que clicaram em orçamento e não converteram) com US$ 5–7/dia em Display.", bold_prefix="🔁 ")
add_bullet("A/B test de headline na landing page focando '1 Day Removal • Hassle-Free • Licensed & Insured'.", bold_prefix="🧪 ")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body("Metas mensuráveis para o próximo ciclo", bold=True, color=NAVY, size=11)
add_table_data(
    headers=["Indicador", "Atual (Mar–Abr)", "Meta (Abr–Mai)"],
    rows=[
        ["CPA",            "US$ 56,16",    "≤ US$ 45,00"],
        ["CTR",            "5,11%",        "≥ 6,00%"],
        ["Conversões",     "35",           "≥ 42"],
        ["Investimento",   "US$ 1.965,77", "Manter ≈ US$ 1.860 (US$ 60/dia)"],
    ],
    col_widths=[5.5, 5.5, 5.5],
)

# =====================================================================
# 9. BUDGET REALLOCATION
# =====================================================================
add_h2("9. Realocação de Verba — Mantendo US$ 60/dia")

add_table_data(
    headers=["Campanha", "Atual", "Sugestão p/ Maio", "Justificativa"],
    rows=[
        ["[Search] Floor Removal — Phone Leads",
         "US$ 30/dia", "US$ 30/dia", "Mantém — campanha principal (CPA US$ 42,63 trimestre)"],
        ["[Search] Specific Services — Floor Demo",
         "US$ 30/dia", "US$ 30/dia", "Mantém — testar Target CPA US$ 45"],
        ["[Search] Brand — Sunrise Floor Removal",
         "PAUSADA",    "PAUSADA",   "Sem mudança até estabilizar CPA não-brand"],
        ["[Display] Remarketing — visitantes",
         "—",          "≈ US$ 5/dia (extra)", "Reativar lista (≥1.000 usuários atingido)"],
    ],
    col_widths=[5.5, 2.5, 3.0, 5.0],
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    "Importante",
    "A campanha de Brand segue pausada para concentrar verba nas campanhas de intenção direta. "
    "Voltará a entrar em pauta quando o CPA não-brand baixar para a faixa de US$ 40–45 e houver "
    "espaço de orçamento. A reativação leve do Remarketing visa recuperar visitantes do site que "
    "clicaram em orçamento mas não converteram, com risco baixo de canibalizar."
)

# =====================================================================
# 10. CLIENT ALIGNMENT
# =====================================================================
add_h2("10. Alinhamento com o Cliente")

add_body(
    "Esse mês fechou abaixo da média do trimestre, mas dentro do padrão de oscilação que "
    "esperamos. O ciclo de melhoria já está mapeado e começa a entrar em execução em Maio:",
    color=GRAY
)

add_bullet("Reforçar a barreira de negativas (concorrentes/varejo) — economiza ~US$ 60–80/mês em tráfego frio.", bold_prefix="Plano em execução — ")
add_bullet("Cortar exposição em condados que não convertem em 91 dias (Pasco, Sumter, Volusia) — libera verba para os condados que entregam.", bold_prefix="Plano em execução — ")
add_bullet("Calibrar Target CPA gradualmente (de US$ 50 → US$ 45) sem comprometer volume.", bold_prefix="Plano em execução — ")
add_bullet("Subir 4 novos anúncios com ângulos de prova social e autoridade para destravar o CTR.", bold_prefix="Plano em execução — ")

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body(
    "A qualidade dos leads continua sendo a referência principal. As ligações e formulários "
    "vindos do Google Ads seguem com perfil compatível com o serviço — esse é o sinal mais "
    "importante e ele se mantém saudável. As métricas dentro da plataforma são meio para o "
    "fim; o que importa de verdade é o ROI real, que continuamos cruzando via planilha "
    "compartilhada.",
    color=GRAY
)

# =====================================================================
# FOOTER
# =====================================================================
doc.add_paragraph().paragraph_format.space_after = Pt(8)

footer_table = doc.add_table(rows=1, cols=1)
fc = footer_table.cell(0, 0)
shade(fc, "F4F6F9")
set_cell_borders(fc, "D5DBE3", "4")
fc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
fp = fc.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run("Relatório gerado a partir dos dados oficiais do Google Ads (CSV exports da conta) — janela 27/02/2025 a 27/03/2025")
fr.font.size = Pt(9)
fr.font.color.rgb = GRAY
fr.italic = True

# =====================================================================
# SAVE
# =====================================================================
out_path = "relatorio_interno.docx"
doc.save(out_path)
print(f"OK: {out_path}")
